from docx import Document
from docx.api import Document

doc = Document()

doc.add_heading('Hello World!', 1)

p = doc.add_paragraph("This is a simple text")
p.add_run("This text is bold").bold = 1
p.add_run("This text is Italic").italic = 1

doc.add_paragraph("This is item one", style='List Bullet')
doc.add_paragraph("This is item two", style='List Bullet')

table_header = ["Name", "Age", "Job"]

some_data = [
    ["Retroider", "34", 'Pidor'],
    ["Edik", "33", "ne pidor"]
]

table = doc.add_table(rows=1, cols=3)
for row_index in range(3):
    table.rows[0].cells[row_index].text = table_header[row_index]

for name, age, job in some_data:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = age
    cells[2].text = job

doc.add_page_break()

doc.add_paragraph("Fuck")

# doc.add_picture()

# doc.save('test.docx')


from docx.api import Document

document = Document('test.docx')

for p in document.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name == "Title":
        print(p.text)

for table in document.tables:
    print("NEW TABLE")
    for row in table.rows:
        print("|".join([cell.text for cell in row.cells]))

all_text = ""
for p in document.paragraphs:
    all_text += p.text
    all_text += '\n'

print(all_text)


