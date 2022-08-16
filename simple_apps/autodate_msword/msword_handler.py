import datetime

from docx.api import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt

WEEKDAYSMAPPING = ["Понедельник", "Вторник",
                   "Среда", "Четверг", "Пятница",
                   "Суббота", "Воскресенье"]


class Column:
    def __init__(self, document_path):
        self._document = Document(document_path)
        self._datetime_list = []
        self._restricted_dates = []

    def get_column(self, col_name):
        for table in self._document.tables:
            for col in table.columns:
                if col_name == col.cells[0].text or col.cells[1].text:
                    return col

    def add_new_datetime(self, new_datetime):
        self._datetime_list.append(new_datetime)

    def sort_datetime_list(self, sort_array):
        if len(sort_array) > 1:
            left_arr = sort_array[:len(sort_array) // 2]
            right_arr = sort_array[len(sort_array) // 2:]

            self.sort_datetime_list(left_arr)
            self.sort_datetime_list(right_arr)

            left_index, right_index, merged_index = 0, 0, 0

            while left_index < len(left_arr) and right_index < len(right_arr):
                if left_arr[left_index] < right_arr[right_index]:
                    sort_array[merged_index] = left_arr[left_index]
                    left_index += 1
                else:
                    sort_array[merged_index] = right_arr[right_index]
                    right_index += 1
                merged_index += 1

            while left_index < len(left_arr):
                sort_array[merged_index] = left_arr[left_index]
                left_index, merged_index = left_index + 1, merged_index + 1
            while right_index < len(right_arr):
                sort_array[merged_index] = right_arr[right_index]
                right_index, merged_index = right_index + 1, merged_index + 1

            if merged_index == len(self._datetime_list):
                self._datetime_list = sort_array

    def get_list_of_columns(self):
        col_names_list = []
        for table in self._document.tables:
            for column in table.columns:
                if len(column.cells[0].text) < 10:
                    col_names_list.append(column.cells[0].text)
                if len(column.cells[1].text) < 10:
                    col_names_list.append(column.cells[1].text)
        return list(set(col_names_list))

    def get_list_of_datetime(self):
        return self._datetime_list

    def _is_restricted_date(self, date):
        for rdate in self._restricted_dates:
            if date >= rdate[0] and date <= rdate[1]:
                return False
            else:
                continue
        return True

    def set_restricted_dates(self, rest_dates):
        self._restricted_dates = rest_dates

    def get_restricted_date(self, date):
        date += datetime.timedelta(days=7)
        while not self._is_restricted_date(date):
            date += datetime.timedelta(days=7)
        return date

    def column_fill(self, col_name):
        self.sort_datetime_list(self._datetime_list)
        for table in self._document.tables:
            for col in table.columns:
                if col_name == col.cells[0].text or col_name == col.cells[1].text:
                    index_of_date_list = 0
                    for cell in col.cells:
                        if cell.text == "":
                            cell.text = ('.'.join(('%02d' % self._datetime_list[index_of_date_list].day,
                                                   '%02d' % self._datetime_list[index_of_date_list].month)))
                            self._datetime_list[index_of_date_list] = \
                                self.get_restricted_date(self._datetime_list[index_of_date_list])
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            index_of_date_list += 1
                            index_of_date_list %= len(self._datetime_list)
        self._document.save('doc' + str(datetime.date.today()) + ".docx")
